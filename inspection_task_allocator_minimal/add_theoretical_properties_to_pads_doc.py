import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_PADS_adaptive_final_theory_patch.docx"
)
OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_PADS_theoretical_properties_added.docx"
)
BACKUP_OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_PADS_theoretical_properties_added.bak_20260503.docx"
)


def set_paragraph_text(paragraph, text):
    paragraph.text = text


def insert_paragraph_after(paragraph, text="", alignment=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if alignment is not None:
        new_para.alignment = alignment
    if text:
        new_para.add_run(text)
    return new_para


def main():
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_PATH}")

    doc = Document(str(SOURCE_PATH))
    paragraphs = doc.paragraphs

    # Tighten earlier wording so the manuscript avoids over-strong theory terms.
    set_paragraph_text(
        paragraphs[79],
        "为降低固定权重设置带来的场景依赖性，A-RH-PADS 引入自适应响应—代价权衡系数 lambda_t。需要说明的是，lambda_t 并非用于提供最优性保证，而是作为任务调度层的在线反馈调节变量，用于根据当前任务紧急度压力、异常压力和路径压力动态改变序列评分函数中响应收益与运动代价的相对权重。由于 lambda_t 由 Sigmoid 函数映射并被限制在 [lambda_min, lambda_max] 区间内，其取值始终有界，避免了响应项或代价项在评分函数中无限放大。",
    )
    set_paragraph_text(
        paragraphs[80],
        "当高优先级任务或异常任务占比增加时，U_t 和 A_t 增大，使 lambda_t 上升，调度策略更偏向关键任务响应；当候选任务空间分布离散、路径压力 D_t 增大时，lambda_t 下降，使调度策略更重视运动代价控制。因此，该机制形成了一种基于任务状态的在线闭环反馈调节过程。本文不对该机制作控制理论意义上的收敛性或稳定性结论，而是通过实验验证其在不同场景下能够产生可解释的响应—代价偏好变化。",
    )

    # Renumber the existing beam-search subsection to make room for the new theory section.
    set_paragraph_text(paragraphs[81], "1.5.2 波束搜索求解流程与复杂度分析")

    # Insert the new theoretical properties section before the experiment chapter.
    anchor = paragraphs[84]
    heading = insert_paragraph_after(
        anchor,
        "1.6 理论性质分析",
        alignment=paragraphs[81].alignment,
    )
    intro = insert_paragraph_after(
        heading,
        "为增强 PADS Framework 的理论表述，下面从问题复杂性、有限候选空间内最优性、搜索空间性质、有界性以及复杂度角度进行简要分析。需要强调的是，这些性质用于说明方法的求解边界、可解释性与在线近似求解合理性，不构成最优性、收敛性或稳定性结论。",
        alignment=paragraphs[82].alignment,
    )

    p1 = insert_paragraph_after(
        intro,
        "命题 1：单机器人在线动态巡检任务调度问题是 NP-hard。",
        alignment=paragraphs[82].alignment,
    )
    p1_proof = insert_paragraph_after(
        p1,
        "证明：当所有任务优先级、区域风险和异常权重相同，且不考虑异常反馈和响应收益，仅以最小化完成所有巡检点的路径长度为目标时，本文问题退化为给定起点的旅行商路径问题。由于 TSP 是 NP-hard 问题，因此本文问题至少与 TSP 一样困难。基于这一复杂性背景，本文采用滚动时域与 Beam Search 实施在线近似求解。",
        alignment=paragraphs[82].alignment,
    )

    p2 = insert_paragraph_after(
        p1_proof,
        "命题 2：在给定候选任务池、预测时域 H 和序列目标函数 J(Q_t) 的条件下，若 Beam 宽度足以保留每一层所有候选序列，则 RH-PADS / A-RH-PADS 等价于对该有限候选空间进行穷举搜索，因此能够得到该候选空间内的最优任务序列。",
        alignment=paragraphs[82].alignment,
    )
    p2_proof = insert_paragraph_after(
        p2,
        "证明：在固定候选任务池和有限预测时域 H 下，可行任务序列数量是有限的。若 Beam Search 在每一层均不发生剪枝，则所有可行序列都会被保留下来并完成评分。此时，算法与对该有限候选空间进行穷举搜索等价，因此评分最大的序列即为该有限候选空间内的最优解。需要说明的是，这一结论仅针对当前有限候选空间，而不是全局任务空间。",
        alignment=paragraphs[82].alignment,
    )

    p3 = insert_paragraph_after(
        p2_proof,
        "命题 3：在候选池和预测时域固定时，增大 Beam 宽度不会降低搜索空间中的最优目标值。",
        alignment=paragraphs[82].alignment,
    )
    p3_proof = insert_paragraph_after(
        p3,
        "证明：较大的 Beam 宽度保留了较小 Beam 宽度下的候选序列集合，或至少包含一个不小于原集合的候选序列集合，因此可搜索空间不会缩小。由于最优目标值是在该集合上取最大值，故在搜索空间层面其最优评分不会下降。该命题仅说明搜索空间中的最优评分不下降，并不意味着真实执行性能一定单调提高。",
        alignment=paragraphs[82].alignment,
    )

    p4 = insert_paragraph_after(
        p3_proof,
        "命题 4：A-RH-PADS 中自适应权衡系数 lambda_t 始终有界。",
        alignment=paragraphs[82].alignment,
    )
    p4_proof = insert_paragraph_after(
        p4,
        "证明：A-RH-PADS 中有 lambda_t = lambda_min + (lambda_max - lambda_min)σ(z_t)，其中 σ(z_t) 为 Sigmoid 函数，满足 0 < σ(z_t) < 1。因此，必有 lambda_min < lambda_t < lambda_max。该性质保证响应收益项和运动代价项在序列目标函数中的相对权重不会被无限放大，从而保持评分结构的有界性与可解释性。",
        alignment=paragraphs[82].alignment,
    )

    p5 = insert_paragraph_after(
        p4_proof,
        "命题 5：当任务优先级、风险、异常权重、路径代价、转向代价和执行时间代价均归一化到 [0,1]，且预测时域 H 有限时，RH-PADS 和 A-RH-PADS 的序列评分函数均有界。",
        alignment=paragraphs[82].alignment,
    )
    p5_proof = insert_paragraph_after(
        p5,
        "证明：在有限长度 H 的任务序列中，每一步的收益项和代价项均为有界量；折扣因子 gamma 属于 (0,1)，因此折扣求和仍然有界。对于 RH-PADS，序列评分由有限个有界收益项、路径/时间惩罚项和早期覆盖奖励项组成；对于 A-RH-PADS，除上述有界代价项外，lambda_t 也始终有界。因此，两种方法的序列评分函数都由有限个有界项线性组合而成，故整体评分有界。",
        alignment=paragraphs[82].alignment,
    )

    insert_paragraph_after(
        p5_proof,
        "复杂度总结：设候选池大小为 M，预测时域为 H，Beam 宽度为 B，地图节点数为 V。A* 单次路径查询复杂度为 O(V log V)。每轮调度决策复杂度近似为 O(HBMV log V)；若使用路径查询缓存，则序列展开阶段可近似为 O(HBM)。总体来看，PADS Framework 面向的是复杂问题上的在线近似求解，其复杂度高于单步贪心，但在有限候选空间、有限时域和有限束宽约束下仍具有可控的在线计算开销。",
        alignment=paragraphs[82].alignment,
    )

    doc.save(str(OUTPUT_PATH))
    shutil.copy2(OUTPUT_PATH, BACKUP_OUTPUT_PATH)
    print(f"Saved theoretical-properties manuscript to: {OUTPUT_PATH}")
    print(f"Saved backup copy to: {BACKUP_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
