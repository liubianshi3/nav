import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_submission_format_PADS_adaptive_bold_revised.docx"
)
OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_PADS_adaptive_final_theory_patch.docx"
)
BACKUP_OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_PADS_adaptive_final_theory_patch.bak_20260503.docx"
)


def set_paragraph_text(paragraph, text):
    paragraph.text = text


def insert_paragraph_after(paragraph, text="", alignment=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if alignment is not None:
        new_para.alignment = alignment
    if text:
        new_para.add_run(text)
    return new_para


def main():
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_PATH}")

    doc = Document(str(SOURCE_PATH))
    paragraphs = doc.paragraphs

    # Tighten the contribution statement into three framework-oriented points.
    set_paragraph_text(
        paragraphs[17],
        "本文的创新点主要体现在以下三个方面：1）提出单机器人在线动态巡检任务调度建模方法，将关键任务响应、异常反馈和运动代价统一到任务序列决策问题中，避免将问题误写为多机器人任务分配；2）构建由 RH-PADS 与 A-RH-PADS 组成的响应—代价可调 PADS Framework，其中 RH-PADS 强调关键任务优先响应，A-RH-PADS 通过 lambda_t 实现响应收益与运动代价之间的自适应反馈调节；3）构建强基线、消融分析、显著性检验和小车运动学仿真共同支撑的系统性实验分析，用于评估框架在关键任务响应与运动代价之间的权衡能力。",
    )

    # Reuse the blank paragraph after the A-RH-PADS formula block as a small subsection heading.
    set_paragraph_text(paragraphs[78], "1.5.1 自适应权衡系数的有界性与反馈调节分析")
    lambda_p1 = insert_paragraph_after(
        paragraphs[78],
        "为降低固定权重设置带来的场景依赖性，A-RH-PADS 引入自适应响应—代价权衡系数 lambda_t。需要说明的是，lambda_t 并非用于保证全局最优解，而是作为任务调度层的在线反馈调节变量，用于根据当前任务紧急度压力、异常压力和路径压力动态改变序列评分函数中响应收益与运动代价的相对权重。由于 lambda_t 由 Sigmoid 函数映射并被限制在 [lambda_min, lambda_max] 区间内，其取值始终有界，避免了响应项或代价项在评分函数中无限放大。",
        alignment=paragraphs[78].alignment,
    )
    insert_paragraph_after(
        lambda_p1,
        "当高优先级任务或异常任务占比增加时，U_t 和 A_t 增大，使 lambda_t 上升，调度策略更偏向关键任务响应；当候选任务空间分布离散、路径压力 D_t 增大时，lambda_t 下降，使调度策略更重视运动代价控制。因此，该机制形成了一种基于任务状态的在线闭环反馈调节过程。本文并不声称该机制具有严格控制理论意义上的收敛性保证，而是通过实验验证其在不同场景下能够产生可解释的响应—代价偏好变化。",
        alignment=paragraphs[78].alignment,
    )

    # Add the motion-cost interface explanation near the end of the kinematic subsection.
    motion_p1 = insert_paragraph_after(
        paragraphs[191],
        "运动代价接口说明：虽然本文实验并未涉及 Gazebo、Nav2 或真实机器人平台，但所提出的 PADS Framework 并不局限于纯几何路径长度。其运动代价项可以接收来自底层运动模型或导航系统的多种代价信息，包括路径长度、预计执行时间、转弯次数、航向变化量和能耗估计等。本文在小车运动学仿真中进一步引入速度、角速度和路径跟踪过程，并记录车辆轨迹长度、执行时间和航向变化量，用于验证任务调度策略在简化运动约束下的表现。",
        alignment=paragraphs[191].alignment,
    )
    insert_paragraph_after(
        motion_p1,
        "该实验的意义在于，PADS 并非仅根据巡检点之间的静态距离排序，而是能够将运动执行层反馈的代价量纳入任务序列决策。实验结果表明，RH-PADS-L 在高优先级任务响应方面仍具有优势，但其车辆轨迹长度和执行时间高于路径效率优先方法，说明关键任务响应与运动代价之间的权衡在引入速度和转向约束后依然存在。后续在 Gazebo、Nav2 或真实机器人平台中，可将底层导航器输出的预计到达时间、局部避障代价、姿态调整耗时和能耗估计进一步接入 PADS 的运动代价项。",
        alignment=paragraphs[191].alignment,
    )

    doc.save(str(OUTPUT_PATH))
    shutil.copy2(OUTPUT_PATH, BACKUP_OUTPUT_PATH)
    print(f"Saved patched manuscript to: {OUTPUT_PATH}")
    print(f"Saved backup copy to: {BACKUP_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
