import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


SOURCE_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_RH_v2_refs_cn_figs_submission_format.docx"
)
OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_submission_format_PADS_adaptive_bold_revised.docx"
)
BACKUP_OUTPUT_PATH = Path(
    "/home/dell/下載/AROC_quadruped_inspection_manuscript_submission_format_PADS_adaptive_bold_revised.bak_20260503.docx"
)


def set_paragraph_text(paragraph, text):
    paragraph.text = text


def replace_text(text):
    replacements = [
        ("RH-Proposed-v2", "RH-PADS"),
        ("RH-v2-Light", "RH-PADS-L"),
        ("RH-v2-Full", "RH-PADS"),
        ("RH-v2-Medium", "RH-PADS-Medium"),
        ("RH-v2-Fast", "RH-PADS-Fast"),
        ("RH-v2", "RH-PADS"),
        ("Proposed-Balanced", "Greedy-PADS"),
        ("DRF-RH", "DRF-PADS"),
        ("动态任务分配", "在线动态任务调度"),
        ("巡检任务分配", "巡检任务调度"),
        ("任务分配层", "任务调度层"),
        ("任务分配方法", "任务调度方法"),
        ("任务分配机制", "任务调度机制"),
        ("任务分配器", "任务调度器"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def insert_paragraph_after(paragraph, text="", alignment=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if alignment is not None:
        new_para.alignment = alignment
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_after_table(table, text="", alignment=None):
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_para = Paragraph(new_p, table._parent)
    if alignment is not None:
        new_para.alignment = alignment
    if text:
        new_para.add_run(text)
    return new_para


def insert_table_after(doc, paragraph, rows, cols, style=None):
    table = doc.add_table(rows=rows, cols=cols)
    if style is not None:
        table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def fill_table(table, data):
    for i, row_data in enumerate(data):
        for j, value in enumerate(row_data):
            table.cell(i, j).text = value


def main():
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_PATH}")

    doc = Document(str(SOURCE_PATH))
    paragraphs = doc.paragraphs

    # Title and abstract front matter.
    set_paragraph_text(
        paragraphs[2],
        "面向关键任务响应与路径代价权衡的单机器人在线动态巡检任务调度框架",
    )
    set_paragraph_text(
        paragraphs[5],
        "摘 要：针对单机器人巡检过程中高优先级任务与异常任务需要快速响应，而距离优先策略容易忽略任务紧急度的问题，本文研究单机器人在线动态巡检任务调度问题。为兼顾关键任务响应与路径代价控制，提出一种优先级感知动态调度框架 PADS。该框架包含响应优先模式 RH-PADS 和自适应扩展模式 A-RH-PADS：前者通过任务紧急度、路径代价、有限预测窗口、波束搜索与滚动执行机制强化关键任务早期响应；后者进一步根据任务紧急度压力、异常压力和路径压力动态调节响应收益与运动代价之间的权衡系数 lambda_t。实验覆盖二维随机栅格、结构化地图、异常反馈、消融分析、强基线对比、统计检验以及小车运动学仿真。结果表明，RH-PADS 在高优先级任务平均响应时间方面表现最好；A-RH-PADS 能进一步降低路径长度和总巡检时间，但会牺牲部分高优先级响应能力，体现出关键任务响应与运动代价之间的可调权衡。小车运动学仿真进一步说明，这一权衡在带速度和转向约束的简化执行模型下仍然存在。本文工作聚焦单机器人巡检任务调度层，不涉及 Gazebo、Nav2 或真实机器人实验。",
    )
    set_paragraph_text(
        paragraphs[6],
        "关键词：单机器人；在线动态巡检任务调度；滚动时域；优先级感知动态调度框架；异常反馈；响应—代价权衡",
    )
    set_paragraph_text(
        paragraphs[8],
        "Priority-Aware Dynamic Scheduling Framework for Online Dynamic Inspection Task Scheduling of a Single Robot",
    )
    set_paragraph_text(
        paragraphs[10],
        "Abstract: This paper studies online dynamic inspection task scheduling for a single robot, where distance-oriented strategies often delay the response to high-priority and abnormal tasks. To balance critical-task response and motion-cost control, a Priority-Aware Dynamic Scheduling framework, denoted as PADS, is proposed. The framework includes a response-priority mode RH-PADS and an adaptive extension mode A-RH-PADS. RH-PADS combines task urgency modeling, path-cost modeling, a finite prediction horizon, beam search and rolling execution to evaluate candidate task sequences. A-RH-PADS further introduces an adaptive response-cost trade-off coefficient lambda_t, which is adjusted online according to urgency pressure, abnormal pressure and path pressure. Experiments on random grid maps, structured maps, abnormal-feedback scenarios, ablation settings, strong baselines, statistical tests and a kinematic vehicle simulator show that RH-PADS achieves the best high-priority response performance, whereas A-RH-PADS reduces path length and total inspection time at the cost of slower high-priority response. The kinematic simulation further confirms that this trade-off remains under simplified speed and turning constraints. The study focuses on the task-scheduling layer only and does not claim Gazebo, Nav2 or real-robot validation.",
    )
    set_paragraph_text(
        paragraphs[11],
        "Key words: single robot; online dynamic inspection task scheduling; receding horizon; priority-aware dynamic scheduling framework; abnormal feedback; response-cost trade-off",
    )

    # Introduction.
    set_paragraph_text(paragraphs[13], "随着移动机器人、自主感知和智能巡检技术的发展，单机器人平台在工业巡检、安防巡逻、设备状态监测和复杂环境探测等场景中得到持续应用。对于此类应用，机器人通常需要在同一张地图上连续访问多个巡检点，而不同任务在重要程度、风险等级和异常反馈方面存在明显差异。")
    set_paragraph_text(paragraphs[14], "在单机器人巡检过程中，距离优先或局部路径代价优先策略虽然能够降低一部分移动开销，但往往会推迟关键设备、高风险区域和异常任务的检测时刻，从而削弱巡检系统的响应价值。因此，问题的关键不只是“走得更省”，而是如何在有限运动代价下更早响应关键任务。")
    set_paragraph_text(paragraphs[15], "本文研究的是单机器人在线动态巡检任务调度问题，即在给定环境地图、机器人当前位置、任务优先级、区域风险和异常反馈的条件下，在线决定机器人后续巡检任务序列。该问题不同于多机器人任务分配，不涉及多机器人之间的资源竞争或协作分配，而更接近带优先级、异常反馈和运动代价约束的在线任务序列优化问题。")
    set_paragraph_text(paragraphs[16], "为此，本文提出优先级感知动态调度框架 PADS（Priority-Aware Dynamic Scheduling Framework）。该框架包含两个模式：RH-PADS（Receding-Horizon Priority-Aware Dynamic Scheduler，响应优先滚动时域动态调度方法）作为本文主方法，强调高优先级任务和异常任务的快速响应；A-RH-PADS（Adaptive Receding-Horizon Priority-Aware Dynamic Scheduler，自适应响应—代价权衡调度方法）作为自适应扩展模式，用于根据任务状态动态调节响应收益与运动代价之间的权衡。")
    set_paragraph_text(paragraphs[17], "本文的主要贡献如下：1）提出单机器人在线动态巡检任务调度建模方法，将关键任务响应、异常反馈和运动代价统一到任务序列决策问题中，避免将问题误写为多机器人任务分配；2）提出响应优先滚动时域调度方法 RH-PADS，通过有限预测窗口与波束搜索缓解单步贪心策略的短视性；3）提出自适应响应—代价权衡扩展方法 A-RH-PADS，通过任务紧急度压力、异常压力和路径压力动态调整 lambda_t，为不同巡检场景提供可调的响应—代价机制；4）构建随机栅格、结构化地图、异常反馈、消融分析、强基线对比、显著性检验和小车运动学仿真的系统化验证体系，用于评估 PADS Framework 在关键任务响应与运动代价之间的权衡能力。")

    # Methods.
    set_paragraph_text(paragraphs[18], "1 PADS Framework 方法设计")
    set_paragraph_text(paragraphs[19], "1.1 问题定义")
    set_paragraph_text(paragraphs[20], "本文研究单机器人在二维栅格地图中的在线动态巡检任务调度问题。机器人在每个决策时刻基于当前位置、未完成任务集合及其属性，在线决定后续巡检任务序列。本文关注任务调度层决策，而不涉及底层导航控制、步态控制或真实平台执行系统。")
    set_paragraph_text(paragraphs[21], "设待执行巡检任务集合为：")
    set_paragraph_text(paragraphs[22], "T = {τ1, τ2, ..., τn}")
    set_paragraph_text(paragraphs[23], "其中，第 i 个任务定义为：")
    set_paragraph_text(paragraphs[24], "τi = (xi, pi, ri, ai)")
    set_paragraph_text(paragraphs[25], "式中，xi 表示任务位置，pi 表示任务优先级，ri 表示区域风险，ai 表示异常反馈权重。机器人在时刻 t 的状态记为 s_t = (x_t, y_t)，预测时域内待评估的候选任务序列记为 Q_t = (τq1, τq2, ..., τqH)。")
    set_paragraph_text(paragraphs[26], "在每个决策时刻，系统从有限候选任务池中构建长度为 H 的序列，仅执行得分最高序列的首个任务；任务完成后根据机器人位置、剩余任务和异常反馈滚动更新。因而，本文问题本质上是带优先级、异常反馈和运动代价约束的在线任务序列决策问题。")
    set_paragraph_text(paragraphs[27], "1.2 任务紧急度与运动代价建模")
    set_paragraph_text(paragraphs[28], "PADS Framework 从任务紧急度和运动代价两个层面刻画调度目标。路径代价仍由 A* 负责估计，但 A* 仅提供可达性、路径长度和转向信息；真正的决策对象是任务序列。")
    set_paragraph_text(paragraphs[29], "在 A-RH-PADS 中，单任务紧急度定义为：")
    set_paragraph_text(paragraphs[30], "ui(t) = clip(0.45pi + 0.35ri + 0.20ai(t), 0, 1)")
    set_paragraph_text(paragraphs[31], "该定义将任务优先级、区域风险和异常反馈聚合为统一紧急度，不再把大量人工权重直接暴露为论文主创新。ui(t) 越大，表示任务越值得被优先完成。")
    set_paragraph_text(paragraphs[32], "在 RH-PADS 的响应优先模式下，仍保留基础单任务效用：")
    set_paragraph_text(paragraphs[33], "Ui(t) = wp pi + wr ri + wa ai(t) - wd di - wc ci - we ei")
    set_paragraph_text(paragraphs[34], "其中，di 为归一化路径距离代价，ci 为路径复杂度代价，ei 为简化能耗代价，所有收益与代价项均归一化到 [0,1]。路径复杂度仍由路径长度、转弯次数和障碍邻近量共同估计。")
    set_paragraph_text(paragraphs[35], "路径复杂度代价可写为：")
    set_paragraph_text(paragraphs[36], "ci = (No + μNt) / (Li + ε)")
    set_paragraph_text(paragraphs[37], "简化能耗项写为 ei = 0.6di + 0.3ci + 0.1Hi，在本文二维栅格调度实验中取 Hi = 0。上述建模为后续序列级评分提供统一的收益—代价基础。")
    set_paragraph_text(paragraphs[38], "1.3 RH-PADS 响应优先滚动时域调度")
    set_paragraph_text(paragraphs[39], "Greedy-PADS 属于单步贪心策略，每轮只依据当前任务评分选择下一任务，容易忽略当前选择对后续路径和关键任务完成顺序的影响。为缓解这一短视性，本文构建响应优先模式 RH-PADS。")
    set_paragraph_text(paragraphs[40], "RH-PADS 在每个决策时刻基于有限预测时域 H、候选池规模 M 和束宽 B，利用 beam search 对未来任务序列进行近似评估。与只看单步评分不同，RH-PADS 在序列层同时考虑累计路径、累计时间和关键任务的早期完成次序。")
    set_paragraph_text(paragraphs[41], "为避免仅按当前分值截断候选任务，RH-PADS 候选池同时保留基础效用较高、路径较短、优先级较高和异常权重较高的候选任务。该设计使算法能够在响应优先前提下，兼顾后续路径代价控制。")
    set_paragraph_text(paragraphs[42], "当最优候选序列确定后，系统只执行该序列的首个任务，然后根据机器人新位置、剩余任务和异常反馈进入下一轮规划。因此，RH-PADS 属于在线滚动时域调度方法，而不是离线固定序列方法。")
    set_paragraph_text(paragraphs[44], "图 1 PADS Framework 总体框架")
    set_paragraph_text(paragraphs[46], "1.4 RH-PADS 序列级目标函数")
    set_paragraph_text(paragraphs[47], "设当前候选任务序列为：")
    set_paragraph_text(paragraphs[48], "Q_t = (τq1, τq2, ..., τqH)")
    set_paragraph_text(paragraphs[49], "RH-PADS 的序列级目标函数写为：")
    set_paragraph_text(paragraphs[50], "J_RH(Q_t) = Σγ^(k-1)U_qk(t) - λ1C_path(Q_t) - λ2C_time(Q_t) - λ3C_priority(Q_t) + λ4B_early(Q_t)")
    set_paragraph_text(paragraphs[51], "其中，C_path(Q_t) 为累计路径惩罚，C_time(Q_t) 为累计执行时间惩罚，C_priority(Q_t) 为优先级加权完成时间惩罚，B_early(Q_t) 为早期高优任务覆盖奖励，γ 为折扣因子。")
    set_paragraph_text(paragraphs[52], "序列中第 k 个任务的预计移动时间为：")
    set_paragraph_text(paragraphs[53], "Tk_move = Lk / v，    Tk = Tk_move + Tins")
    set_paragraph_text(paragraphs[54], "其中，v 为机器人平均速度，Tins 为单任务巡检时间。基于累计完成时刻 Ck，RH-PADS 对序列前段关键任务的覆盖情况和完成次序进行联合评估。")
    set_paragraph_text(paragraphs[55], "相应地，可写出：")
    set_paragraph_text(paragraphs[56], "C_path(Q_t) = ΣLk / Lscale，    C_time(Q_t) = ΣTk / Tscale")
    set_paragraph_text(paragraphs[57], "C_priority(Q_t) = ΣpkCk / Tscale，    B_early(Q_t) = Nhigh^K / K")
    set_paragraph_text(paragraphs[58], "因此，RH-PADS 的核心优势不在于单轮路径最省，而在于通过有限候选池上的近似启发式序列搜索，使关键任务响应与移动代价控制被同时纳入当前决策。")
    set_paragraph_text(paragraphs[59], "")
    set_paragraph_text(paragraphs[60], "")
    set_paragraph_text(paragraphs[61], "")
    set_paragraph_text(paragraphs[62], "")
    set_paragraph_text(paragraphs[63], "")
    set_paragraph_text(paragraphs[64], "")
    set_paragraph_text(paragraphs[65], "")
    set_paragraph_text(paragraphs[66], "")
    set_paragraph_text(paragraphs[67], "1.5 A-RH-PADS 自适应响应—代价权衡扩展")
    set_paragraph_text(paragraphs[68], "A-RH-PADS 是 PADS Framework 的自适应扩展模式，用于在不同巡检状态下动态调节响应收益与运动代价之间的权衡。它不是 RH-PADS 的全面替代，而是代价约束增强场景下的补充模式。")
    set_paragraph_text(paragraphs[69], "当前任务紧急度压力定义为：")
    set_paragraph_text(paragraphs[70], "U_t = 0.6 max_i ui(t) + 0.4 mean(top-k ui(t))")
    set_paragraph_text(paragraphs[71], "异常压力与路径压力分别定义为：")
    set_paragraph_text(paragraphs[72], "A_t = max_i ai(t)，    D_t = mean_i d̄i(t)")
    set_paragraph_text(paragraphs[73], "自适应权衡系数写为：")
    set_paragraph_text(paragraphs[74], "lambda_t = lambda_min + (lambda_max - lambda_min)σ(k0 + kuU_t + kaA_t - kdD_t)，其中 lambda_min = 0.25，lambda_max = 0.85。lambda_t 越大表示越偏响应优先，越小表示越偏运动代价控制。")
    set_paragraph_text(paragraphs[75], "A-RH-PADS 的序列目标写为：")
    set_paragraph_text(paragraphs[76], "J_A(Q_t) = lambda_t R(Q_t) - (1 - lambda_t) C(Q_t)")
    set_paragraph_text(paragraphs[77], "其中，R(Q_t) = Σ discount^(k-1)u_qk(t) / (1 + finish_time_norm_qk)，C(Q_t) = 0.70C_path + 0.20C_turn + 0.10C_time。该目标使 A-RH-PADS 更适合用于响应—代价需要自适应调整的巡检场景。")
    set_paragraph_text(paragraphs[79], "1.6 波束搜索求解流程与复杂度分析")

    # Experiment structure and captions.
    set_paragraph_text(paragraphs[82], "为验证 PADS Framework 的有效性，本文构建二维栅格仿真实验环境。实验目的在于评估任务调度层算法在不同任务属性、路径代价和异常反馈条件下的表现，而非验证 ROS2、Nav2、Gazebo 或真实机器人平台中的导航执行性能。底层路径代价统一由 A* 计算，本文方法负责在线巡检任务序列决策。")
    set_paragraph_text(paragraphs[83], "实验包括随机栅格主实验、多规模泛化实验、参数效率实验、异常反馈实验、消融实验、权重敏感性实验、强基线对比、结构化地图泛化、小车运动学仿真以及 A-RH-PADS 自适应扩展实验。主实验采用 30×30 栅格地图、20 个巡检任务、障碍物比例 0.2，并在 30 个随机种子下重复运行。除二维栅格统计实验外，本文还构建了差速小车运动学仿真，用于观察任务调度结果在带速度和转向约束的简化执行模型下的表现。该实验同样不涉及 Gazebo、Nav2 或真实机器人平台。")
    set_paragraph_text(paragraphs[84], "对比方法包括 NNF、AStarOnly、Greedy-PADS 和 RH-PADS；在扩展实验中，进一步引入 Priority-Greedy、Deadline-Greedy、TSP-2opt、DRF-PADS-Full、DRF-PADS-Light、A-RH-PADS 和 A-RH-PADS-L。其中，RH-PADS 是本文主方法，A-RH-PADS 是用于动态调节响应—代价权衡的扩展模式，DRF-PADS 则被定位为探索性增强版本。")
    set_paragraph_text(paragraphs[85], "评价指标包括总路径长度、总巡检时间、高优先级任务平均响应时间、优先级加权完成时间、前 5 个任务中高优先级任务比例、异常任务优先处理率、异常平均响应时间、lambda_t 统计量以及算法运行时间。其中，路径、时间和响应类指标越小越好；高优先级覆盖率和异常任务优先处理率越大越好。")
    set_paragraph_text(paragraphs[86], "2.2 随机栅格主实验")
    set_paragraph_text(paragraphs[87], "表 1 RH-PADS 主实验性能对比")
    set_paragraph_text(paragraphs[89], "表 2 RH-PADS 相比 Greedy-PADS 和 AStarOnly 的主实验指标变化率")
    set_paragraph_text(paragraphs[91], "由表 1 和表 2 可见，RH-PADS 相比 Greedy-PADS 在路径代价和关键任务响应两方面均有改进。具体而言，RH-PADS 的总路径长度由 236.63 降低至 220.27，总巡检时间由 494.06 s 降低至 466.78 s；同时，高优先级任务平均响应时间由 170.82 s 降低至 130.12 s，优先级加权完成时间同步下降。这说明响应优先滚动时域调度能够缓解单步贪心策略的短视性。")
    set_paragraph_text(paragraphs[92], "与 AStarOnly 相比，RH-PADS 的总路径长度和总巡检时间仍然更高，因此 PADS Framework 不是面向最低路径代价的框架。但在关键任务响应方面，RH-PADS 的高优先级任务平均响应时间明显更低，体现出“关键任务优先”而非“距离优先”的核心定位。")
    set_paragraph_text(paragraphs[93], "总体来看，随机栅格主实验表明 RH-PADS 通过有限时域序列评分，在保持在线决策能力的同时改善了关键任务早期响应，并相较 Greedy-PADS 获得了更平衡的路径—响应折中。")
    set_paragraph_text(paragraphs[95], "图 2 随机栅格主实验中不同方法的性能对比")
    set_paragraph_text(paragraphs[97], "2.3 多规模泛化实验")
    set_paragraph_text(paragraphs[98], "表 3 多规模泛化实验总体结果")
    set_paragraph_text(paragraphs[100], "多规模泛化实验覆盖不同地图规模、任务数量和障碍物比例。从总体均值看，RH-PADS 的高优先级任务平均响应时间为 146.65 s，明显低于 NNF、AStarOnly 和 Greedy-PADS，说明响应优先滚动时域调度在更广泛配置下仍能保持关键任务响应优势。")
    set_paragraph_text(paragraphs[101], "在路径与时间方面，RH-PADS 的总路径长度和总巡检时间均低于 Greedy-PADS，表明序列级评估不仅提高了响应能力，也在一定程度上缓解了跨区域绕行问题。")
    set_paragraph_text(paragraphs[102], "同时，RH-PADS 在总路径长度和总巡检时间上仍高于 NNF 和 AStarOnly，这进一步说明该方法并不是路径效率优先方法，而是围绕关键任务响应与移动代价之间的折中进行设计。")
    set_paragraph_text(paragraphs[104], "图 3 多规模场景下各方法总体性能对比")
    set_paragraph_text(paragraphs[106], "2.4 参数效率实验")
    set_paragraph_text(paragraphs[107], "表 4 RH-PADS 参数效率实验结果")
    set_paragraph_text(paragraphs[109], "参数效率实验表明，RH-PADS 采用较大的预测步长和束宽时，高优先级任务平均响应时间最低，说明更充分的序列展开有利于响应优先决策。")
    set_paragraph_text(paragraphs[110], "随着参数规模减小，算法运行时间逐步下降。RH-PADS-L 的运行时间低于完整版本，而高优先级任务平均响应时间和早期高优覆盖率只出现有限退化，说明其在性能与计算开销之间取得了较好的折中。")
    set_paragraph_text(paragraphs[111], "因此，RH-PADS 可作为效果优先版本展示响应优势，RH-PADS-L 则可作为轻量化版本用于计算资源受限场景。表 4 中的运行时间基于带 A* 路径查询缓存的设置，仅用于比较不同参数组的相对计算开销。")
    set_paragraph_text(paragraphs[113], "图 4 RH-PADS 参数设置对性能与开销的影响")
    set_paragraph_text(paragraphs[116], "图 5 RH-PADS 参数效率权衡关系")
    set_paragraph_text(paragraphs[119], "2.5 异常反馈实验")
    set_paragraph_text(paragraphs[120], "表 5 RH-PADS 异常反馈实验结果")
    set_paragraph_text(paragraphs[122], "异常反馈实验表明，异常触发后，RH-PADS 的异常平均响应时间和异常任务优先处理率均优于 AStarOnly，并整体优于 Greedy-PADS，说明异常反馈进入序列级目标函数后能够促进异常任务更早进入后续执行序列。")
    set_paragraph_text(paragraphs[123], "RH-PADS-L 的异常平均响应时间同样低于 Greedy-PADS 和 AStarOnly，表明轻量化版本仍保持较好的异常响应能力，但其异常任务优先处理率略低于 RH-PADS，体现了轻量化搜索带来的折中。")
    set_paragraph_text(paragraphs[124], "在路径和总时间方面，RH-PADS 与 RH-PADS-L 仍低于 Greedy-PADS，表明滚动时域序列评分在异常反馈场景下依然能够兼顾路径代价控制与关键任务响应。")
    set_paragraph_text(paragraphs[125], "需要进一步指出的是，作为探索性扩展的 DRF-PADS 在异常任务响应上并未稳定超过 RH-PADS。与 RH-PADS 相比，DRF-PADS-Full 的异常平均响应时间更高、异常优先处理率更低，因此动态异常风险场在当前设置下更适合作为后续优化方向，而非本文主线结论。")
    set_paragraph_text(paragraphs[127], "图 6 异常反馈实验结果对比")
    set_paragraph_text(paragraphs[130], "2.6 消融实验")
    set_paragraph_text(paragraphs[131], "表 6 Greedy-PADS 权重项消融实验结果")
    set_paragraph_text(paragraphs[133], "消融实验结果表明，去除 priority 项后，高优先级任务平均响应时间和早期高优覆盖率均明显退化，说明任务优先级对关键任务的早期响应具有决定性作用。")
    set_paragraph_text(paragraphs[134], "去除 risk 项后，高风险任务响应时间升高；去除 abnormal_weight 项后，异常任务优先处理率大幅下降且异常平均响应时间上升，说明风险和异常反馈分别主导高风险区域任务和异常任务的调度顺序。")
    set_paragraph_text(paragraphs[135], "需要如实指出的是，去除 complexity 项后部分指标反而改善，说明当前路径复杂度项尚未表现出稳定的正向收益，其尺度设计和统计方式仍有进一步优化空间。")
    set_paragraph_text(paragraphs[136], "相比之下，去除 energy 项后的变化相对有限，说明在当前二维栅格和简化能耗估计条件下，能耗项更多体现为可扩展接口，而非主导性决定因素。")
    set_paragraph_text(paragraphs[138], "图 7 消融实验结果对比")
    set_paragraph_text(paragraphs[141], "2.7 权重敏感性实验")
    set_paragraph_text(paragraphs[142], "表 7 权重敏感性实验结果")
    set_paragraph_text(paragraphs[144], "权重敏感性实验表明，当路径代价被过度强调时，总路径长度和总巡检时间下降，但高优先级任务平均响应时间明显上升，说明过度追求路径效率会削弱关键任务响应能力。")
    set_paragraph_text(paragraphs[145], "当任务优先级被过度强调时，高优先级响应进一步改善，但路径长度和总时间均明显增加，说明收益优先策略会引起更高的运动代价。")
    set_paragraph_text(paragraphs[146], "Balanced 权重在路径效率与关键任务响应之间取得了相对均衡的折中，因此本文将其作为 Greedy-PADS 和 RH-PADS 的基础权重设置。")
    set_paragraph_text(paragraphs[148], "图 8 权重敏感性实验结果对比")
    set_paragraph_text(paragraphs[149], "2.8 强基线对比与统计分析")
    set_paragraph_text(paragraphs[150], "为进一步降低“本文方法仅相对弱基线有效”的质疑，本文在 RH-PADS 之外补充 Priority-Greedy、Deadline-Greedy 和 TSP-2opt 三类更强对比方法。与此同时，引入 DRF-PADS-Full 和 DRF-PADS-Light 作为 PADS Framework 的探索性增强版本，用于验证更强序列目标函数与更强候选池设计的潜在收益。")
    set_paragraph_text(paragraphs[151], "主实验结果表明，DRF-PADS-Full 相比 RH-PADS 的总路径长度和总巡检时间略有增加，但高优先级任务平均响应时间更低，说明其主要收益集中在更激进的关键任务早期响应，而不是更低的路径代价。")
    set_paragraph_text(paragraphs[152], "与 Priority-Greedy 和 TSP-2opt 相比，DRF-PADS-Full 同样体现出典型的响应—代价折中：前者路径与时间更高，后者高优先级响应明显更弱。该结果进一步说明本文框架并非以最低路径代价为唯一目标，而是围绕关键任务响应与运动代价控制展开。")
    set_paragraph_text(paragraphs[153], "表 8 DRF-PADS 主实验强基线对比结果")
    set_paragraph_text(paragraphs[154], "与 Deadline-Greedy 相比，DRF-PADS-Full 的路径和时间代价更低，高优先级响应也更快；与 TSP-2opt 相比，DRF-PADS-Full 的路径和时间代价更高，但关键任务响应明显更强。由此可见，路径效率优先和响应优先之间存在稳定的结构性权衡。")
    set_paragraph_text(paragraphs[155], "统计检验摘要见表 9。结果显示，DRF-PADS-Full 相比 RH-PADS 在高优先级任务平均响应时间和前 5 个任务高优先级比例上均达到统计显著改善，p 值分别为 0.0010 和 0.0014；而总路径长度差异不显著，p = 0.4976。该结果表明，DRF-PADS 的增益主要集中在响应指标，而不是稳定降低路径代价。")
    set_paragraph_text(paragraphs[160], "图 9 DRF-PADS 主实验强基线对比图")
    set_paragraph_text(paragraphs[161], "表 9 显著性检验摘要")
    set_paragraph_text(paragraphs[163], "2.9 DRF-PADS 内部结构消融实验")
    set_paragraph_text(paragraphs[164], "为验证 DRF-PADS 的改进并非简单调参，本文对其内部结构进行消融分析，结果见表 10 和图 10。消融对象包括多步滚动预测、序列级路径时间惩罚、优先级完成时间惩罚、deadline violation penalty、Top-K 早期覆盖奖励、混合候选池以及动态异常风险场。")
    set_paragraph_text(paragraphs[165], "实验结果表明，DRF-PADS-BaseScoreOnly 明显退化，说明不能只累加基础评分；DRF-PADS-NoPathTimePenalty 的总路径长度和总巡检时间明显变差，说明路径/时间序列惩罚对于控制跨区域绕行是必要的；DRF-PADS-NoTopKBonus 的高优先级响应时间和 Top5 指标明显退化，说明 Top-K 早期覆盖奖励对关键任务前期响应具有直接作用。")
    set_paragraph_text(paragraphs[166], "此外，DRF-PADS-NoPriorityCompletion 也出现一定退化，说明优先级完成时间惩罚对关键任务完成顺序具有贡献。需要如实指出的是，DRF-PADS-NoDeadlinePenalty 与 DRF-PADS-Full 在当前主实验设置下几乎完全一致，说明当前 deadline penalty 尚未被有效触发；DRF-PADS-NoDynamicRiskField 在无异常消融中的差异也较有限。")
    set_paragraph_text(paragraphs[167], "表 10 DRF-PADS 内部结构消融结果")
    set_paragraph_text(paragraphs[169], "图 10 DRF-PADS 内部消融实验结果对比")
    set_paragraph_text(paragraphs[172], "2.10 结构化地图泛化实验")
    set_paragraph_text(paragraphs[173], "除随机栅格地图外，本文进一步构建 corridor、room_corridor 和 bottleneck 三类结构化地图，以观察方法在走廊、房间走廊和瓶颈拓扑中的泛化表现，结果见表 11 和图 11。")
    set_paragraph_text(paragraphs[174], "在 corridor 场景下，DRF-PADS-Full 相比 RH-PADS 的路径和时间代价更低，高优先级响应也略优；在 room_corridor 场景下，路径与时间略高但高优先级响应略有改善；在 bottleneck 场景下，路径与时间更低而高优先级响应也略优。")
    set_paragraph_text(paragraphs[175], "上述结果表明，结构化地图中 DRF-PADS 并非在所有拓扑上都获得一致优势，但其收益更接近于减弱跨区绕行并维持关键任务早期响应，而不是在全部结构下同时实现更低路径代价和更快响应。")
    set_paragraph_text(paragraphs[176], "表 11 结构化地图实验结果")
    set_paragraph_text(paragraphs[178], "图 11 结构化地图实验结果对比")
    set_paragraph_text(paragraphs[183], "2.11 小车运动学仿真验证")
    set_paragraph_text(paragraphs[184], "为进一步验证所提任务调度方法在带运动约束执行模型下的适用性，本文构建二维差速小车运动学仿真环境。该实验不是 Gazebo、Nav2 或真实机器人实验，而是在二维栅格地图基础上增加小车速度、角速度、路径跟踪和到达判定过程，用于补充评估任务调度结果在简化移动平台执行模型下的表现。")
    set_paragraph_text(paragraphs[185], "实验选取 AStarOnly、Greedy-PADS、RH-PADS-L、TSP-2opt 和 Priority-Greedy 五种方法进行对比，结果见表 12、图 12 和图 13。结果表明，各方法的平均任务完成数均为 19.93，目标到达成功率均为 100%，说明小车运动学仿真框架能够稳定执行不同任务调度策略。")
    set_paragraph_text(paragraphs[186], "表 12 小车运动学仿真实验结果")
    set_paragraph_text(paragraphs[187], "从车辆轨迹长度和执行时间看，TSP-2opt 与 AStarOnly 表现最好，说明路径效率优先策略在运动学仿真中仍具有较低执行代价。相比之下，RH-PADS-L 的车辆轨迹长度和执行时间更高，说明其为了更早处理高价值任务付出了一定运动代价。")
    set_paragraph_text(paragraphs[188], "从高优先级任务平均响应时间看，RH-PADS-L 取得最低值，为 125.70 s，低于 AStarOnly 的 183.09 s、TSP-2opt 的 183.70 s、Greedy-PADS 的 172.78 s 和 Priority-Greedy 的 151.91 s。该结果说明，在考虑小车速度、转向和路径跟踪过程后，RH-PADS-L 仍能保持关键任务优先响应能力。")
    set_paragraph_text(paragraphs[189], "图 12 小车运动学仿真实验结果对比")
    set_paragraph_text(paragraphs[191], "图 13 RH-PADS-L 小车运动学仿真轨迹示例")
    set_paragraph_text(paragraphs[193], "2.13 小结")
    set_paragraph_text(paragraphs[194], "综合本章结果可知，RH-PADS 仍是 PADS Framework 的响应优先主模式。相较单步 Greedy-PADS，RH-PADS 在随机栅格主实验、多规模实验和异常反馈实验中均表现出更好的关键任务响应能力，并在多数情况下改善了路径代价与响应效率之间的折中关系。")
    set_paragraph_text(paragraphs[195], "扩展实验进一步表明，强基线与内部结构消融验证了序列级目标函数、路径时间惩罚和 Top-K 早期覆盖奖励的有效性；但 DRF-PADS 并未在所有指标上稳定优于 RH-PADS，尤其是动态异常风险场在当前设置下尚未带来更好的异常任务响应。")
    set_paragraph_text(paragraphs[196], "因此，PADS Framework 更适用于关键任务和异常任务需要优先处理的单机器人巡检场景。实验也说明滚动时域搜索会带来额外计算开销，实际部署时可根据计算资源选择 RH-PADS 或 RH-PADS-L 等不同配置。")
    set_paragraph_text(paragraphs[197], "3 结论")
    set_paragraph_text(paragraphs[198], "本文提出面向单机器人在线动态巡检任务调度的 PADS Framework，并围绕关键任务响应与路径代价权衡构建了 RH-PADS 与 A-RH-PADS 两种模式。RH-PADS 作为响应优先模式，通过任务紧急度、路径代价、滚动时域序列搜索和波束搜索求解机制，使关键任务完成顺序被纳入当前决策。")
    set_paragraph_text(paragraphs[199], "实验结果表明，RH-PADS 在高优先级任务响应方面表现最好；A-RH-PADS 作为自适应扩展模式，能够降低路径长度和总执行时间，但会牺牲部分高优先级响应能力。这说明关键任务响应和路径代价之间存在明确权衡，而 PADS Framework 可以根据应用偏好在响应优先与代价控制之间进行模式选择。")
    set_paragraph_text(paragraphs[200], "小车运动学仿真进一步说明，上述权衡在带速度、转向和路径跟踪约束的简化执行模型中仍然存在。与此同时，本文也如实发现：A-RH-PADS 并未全面优于 RH-PADS，DRF-PADS 的增强项也尚未形成稳定的异常响应收益，因此这些扩展更适合作为框架内的能力补充，而非替代 RH-PADS 的统一结论。")
    set_paragraph_text(paragraphs[201], "本文工作聚焦任务调度层，不声称 Gazebo、Nav2 或真实机器人实验已经完成。后续研究将面向更高保真仿真和真实平台验证 PADS Framework 与底层导航系统的接口机制，并进一步优化 lambda_t 的启发式设计、异常传播模型和复杂环境下的路径复杂度建模。")

    # Global method naming updates in body and tables before references.
    for i, paragraph in enumerate(doc.paragraphs[:202]):
        if i in {
            2, 5, 6, 8, 10, 11, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24,
            25, 26, 27, 28, 29, 30, 31, 32, 33, 34, 35, 36, 37, 38, 39, 40, 41,
            42, 44, 46, 47, 48, 49, 50, 51, 52, 53, 54, 55, 56, 57, 58, 67, 68,
            69, 70, 71, 72, 73, 74, 75, 76, 77, 79, 82, 83, 84, 85, 86, 87, 89,
            91, 92, 93, 95, 97, 98, 100, 101, 102, 104, 106, 107, 109, 110, 111,
            113, 116, 119, 120, 122, 123, 124, 125, 127, 130, 131, 133, 134, 135,
            136, 138, 141, 142, 144, 145, 146, 148, 149, 150, 151, 152, 153, 154,
            155, 160, 161, 163, 164, 165, 166, 167, 169, 172, 173, 174, 175, 176,
            178, 183, 184, 185, 186, 187, 188, 189, 191, 193, 194, 195, 196, 197,
            198, 199, 200, 201
        }:
            paragraph.text = replace_text(paragraph.text)

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("作者简介："):
            paragraph.text = paragraph.text.replace("任务分配", "任务调度")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    # Update summary tables with new method names where needed.
    # Table 4 naming cleanup.
    doc.tables[4].cell(1, 0).text = "RH-PADS"
    doc.tables[4].cell(2, 0).text = "RH-PADS-Medium"
    doc.tables[4].cell(3, 0).text = "RH-PADS-L"
    doc.tables[4].cell(4, 0).text = "RH-PADS-Fast"
    doc.tables[5].cell(2, 0).text = "Greedy-PADS"
    doc.tables[6].cell(1, 0).text = "Greedy-PADS"

    # Insert complexity analysis paragraphs after paragraph 79.
    complexity_anchor = paragraphs[79]
    complexity_p1 = insert_paragraph_after(
        complexity_anchor,
        "PADS Framework 在每轮决策中首先以有限候选池构建候选任务集合，再利用 beam search 扩展长度为 H 的部分序列，并采用滚动执行方式只提交最优序列的首个任务。异常反馈触发后，仅更新未完成任务的异常权重并重新规划，不对已完成任务回溯。",
        alignment=complexity_anchor.alignment,
    )
    complexity_p2 = insert_paragraph_after(
        complexity_p1,
        "设候选池大小为 M，波束宽度为 B，预测时域长度为 H，地图节点数为 V。A* 单次路径查询复杂度为 O(V log V)，则 RH-PADS / A-RH-PADS 每轮决策复杂度近似为 O(HBMV log V)。若路径查询缓存可用，则序列展开阶段可近似降低为 O(HBM)。",
        alignment=complexity_anchor.alignment,
    )
    insert_paragraph_after(
        complexity_p2,
        "因此，滚动时域方法的复杂度高于单步贪心策略，但通过限制候选池规模、束宽和预测时域，RH-PADS 与 A-RH-PADS 仍可用于在线任务调度。RH-PADS-L 与 A-RH-PADS-L 则是面向计算资源受限场景的轻量版本。",
        alignment=complexity_anchor.alignment,
    )

    # Insert adaptive experiment section before the summary subsection.
    summary_heading = paragraphs[193]
    fig13_caption = paragraphs[191]
    heading_alignment = summary_heading.alignment
    caption_alignment = paragraphs[186].alignment
    normal_alignment = paragraphs[194].alignment
    table_style = doc.tables[12].style

    adaptive_heading = insert_paragraph_after(
        fig13_caption,
        "2.12 A-RH-PADS 自适应响应—代价权衡扩展实验",
        alignment=heading_alignment,
    )
    adaptive_p1 = insert_paragraph_after(
        adaptive_heading,
        "为进一步验证 PADS Framework 的可调权衡能力，本文将 A-RH-PADS 作为自适应扩展模式单独分析。其核心结论是：A-RH-PADS 能降低路径代价和总巡检时间，但高优先级任务响应不如 RH-PADS。因此，A-RH-PADS 不能替代响应优先模式 RH-PADS，而更适合作为代价受限场景下的扩展模式。",
        alignment=normal_alignment,
    )

    caption13 = insert_paragraph_after(
        adaptive_p1,
        "表 13 A-RH-PADS 自适应扩展主实验结果",
        alignment=caption_alignment,
    )
    table13 = insert_table_after(doc, caption13, 3, 6, style=table_style)
    fill_table(
        table13,
        [
            ["方法", "总路径长度", "总巡检时间/s", "高优响应时间/s", "优先级加权完成时间/s", "λ均值"],
            ["RH-PADS", "220.27", "466.78", "130.12", "183.10", "—"],
            ["A-RH-PADS", "207.07", "444.78", "143.00", "178.72", "0.572"],
        ],
    )

    adaptive_p2 = insert_paragraph_after_table(
        table13,
        "主实验结果表明，A-RH-PADS 使总路径长度由 220.27 降至 207.07，总巡检时间由 466.78 s 降至 444.78 s，优先级加权完成时间由 183.10 s 降至 178.72 s；但高优先级任务平均响应时间由 130.12 s 增至 143.00 s。显著性检验表明，路径改善显著，p = 0.00054；高优先级响应变差同样显著，p = 0.00561。该结果说明自适应模式更偏向代价控制，而不是更激进的关键任务优先响应。",
        alignment=normal_alignment,
    )

    caption14 = insert_paragraph_after(
        adaptive_p2,
        "表 14 A-RH-PADS 自适应扩展异常反馈结果",
        alignment=caption_alignment,
    )
    table14 = insert_table_after(doc, caption14, 3, 5, style=table_style)
    fill_table(
        table14,
        [
            ["方法", "异常优先率/%", "异常响应时间/s", "异常前λ", "异常后λ"],
            ["RH-PADS", "52.50", "116.64", "—", "—"],
            ["A-RH-PADS", "54.17", "118.42", "0.618", "0.785"],
        ],
    )

    adaptive_p3 = insert_paragraph_after_table(
        table14,
        "异常反馈实验表明，A-RH-PADS 在异常触发后可将 lambda_t 从 0.618 提升至 0.785，平均变化量为 +0.167，说明自适应机制能够响应异常压力。然而，A-RH-PADS 的异常平均响应时间为 118.42 s，略高于 RH-PADS 的 116.64 s；对应差异不显著，p = 0.732。换言之，lambda_t 的上升能够解释自适应机制的行为变化，但并未让异常响应稳定超过 RH-PADS。",
        alignment=normal_alignment,
    )

    caption15 = insert_paragraph_after(
        adaptive_p3,
        "表 15 A-RH-PADS-L 小车运动学仿真对比结果",
        alignment=caption_alignment,
    )
    table15 = insert_table_after(doc, caption15, 3, 4, style=table_style)
    fill_table(
        table15,
        [
            ["方法", "车辆轨迹长度", "执行时间/s", "高优响应时间/s"],
            ["RH-PADS-L", "211.95", "420.87", "125.70"],
            ["A-RH-PADS-L", "193.74", "382.57", "162.10"],
        ],
    )

    insert_paragraph_after_table(
        table15,
        "在小车运动学仿真中，A-RH-PADS-L 的车辆轨迹长度和执行时间均低于 RH-PADS-L，但高优先级任务平均响应时间由 125.70 s 升至 162.10 s，且响应退化具有统计显著性，p = 3.55×10^-6；相对地，轨迹长度改善也具有统计显著性，p = 6.21×10^-5。该结果进一步说明，自适应模式在带速度和转向约束的执行模型下同样更偏代价控制。",
        alignment=normal_alignment,
    )

    # Save outputs.
    doc.save(str(OUTPUT_PATH))
    shutil.copy2(OUTPUT_PATH, BACKUP_OUTPUT_PATH)
    print(f"Saved revised manuscript to: {OUTPUT_PATH}")
    print(f"Saved backup copy to: {BACKUP_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
